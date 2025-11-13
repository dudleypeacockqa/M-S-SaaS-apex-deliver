"""
Document Export Service
Feature: F-009 Automated Document Generation - Export to PDF/DOCX
"""
import io
import os
import uuid
from pathlib import Path
from typing import Optional, Dict, Any, BinaryIO
from datetime import datetime, UTC

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from app.services.storage_service import get_storage_service
from app.core.config import get_settings


class DocumentExportService:
    """Service for exporting documents to PDF/DOCX"""

    @staticmethod
    async def export_to_pdf(
        content: str,
        organization_id: str,
        document_id: str,
        options: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        Export document content to PDF file.

        Args:
            content: Document content (plain text or HTML)
            organization_id: Organization UUID
            document_id: Document UUID
            options: Export options (margin, font_family, include_cover_page, etc.)

        Returns:
            File path where PDF is stored

        Raises:
            ImportError: If reportlab is not installed
            IOError: If file generation fails
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is not installed. Install it with: pip install reportlab")

        options = options or {}
        margin = options.get("margin", 15)  # Default 15mm margin
        font_family = options.get("font_family", "Helvetica")
        include_cover_page = options.get("include_cover_page", False)

        # Create buffer for PDF
        buffer = io.BytesIO()

        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,  # Use letter size (8.5" x 11")
            rightMargin=margin,
            leftMargin=margin,
            topMargin=margin,
            bottomMargin=margin,
        )

        # Create styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=18,
            textColor="black",
            spaceAfter=30,
            alignment=TA_CENTER,
        )
        body_style = ParagraphStyle(
            "CustomBody",
            parent=styles["Normal"],
            fontSize=11,
            fontName=font_family,
            spaceAfter=12,
            alignment=TA_LEFT,
        )

        # Build PDF content
        story = []

        if include_cover_page:
            # Add cover page
            story.append(Spacer(1, 2 * inch))
            story.append(Paragraph("Document", title_style))
            story.append(Spacer(1, 0.5 * inch))
            story.append(Paragraph(
                f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M:%S UTC')}",
                body_style
            ))
            story.append(PageBreak())

        # Split content into paragraphs and add to PDF
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                # Clean up HTML-like tags if present (simple cleaning)
                para_text = para_text.replace("<p>", "").replace("</p>", "")
                para_text = para_text.replace("<br>", "\n").replace("<br/>", "\n")
                para_text = para_text.replace("&nbsp;", " ")
                
                story.append(Paragraph(para_text, body_style))
                story.append(Spacer(1, 0.2 * inch))

        # Build PDF
        doc.build(story)

        # Get PDF bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()

        # Save to storage
        storage_service = get_storage_service()
        # Generate file key using storage service (requires deal_id, but we don't have one for exports)
        # Use a placeholder deal_id or document_id
        file_key = storage_service.generate_file_key(
            organization_id=organization_id,
            deal_id=document_id,  # Use document_id as deal_id placeholder
            filename=f"{document_id}.pdf",
            user_id=organization_id,  # Use organization_id as user_id placeholder
        )
        
        # Create file stream from bytes
        file_stream = io.BytesIO(pdf_bytes)
        
        # Save file (storage service expects async)
        file_path = await storage_service.save_file(
            file_key=file_key,
            file_stream=file_stream,
            organization_id=organization_id,
        )

        # Return file_key for retrieval, not full path
        return file_key

    @staticmethod
    async def export_to_docx(
        content: str,
        organization_id: str,
        document_id: str,
        options: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        Export document content to DOCX file.

        Args:
            content: Document content (plain text or HTML)
            organization_id: Organization UUID
            document_id: Document UUID
            options: Export options (font_family, font_size, etc.)

        Returns:
            File path where DOCX is stored

        Raises:
            ImportError: If python-docx is not installed
            IOError: If file generation fails
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")

        options = options or {}
        font_family = options.get("font_family", "Calibri")
        font_size = options.get("font_size", 11)
        include_cover_page = options.get("include_cover_page", False)

        # Create DOCX document
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = font_family
        font.size = Pt(font_size)

        if include_cover_page:
            # Add cover page
            title = doc.add_heading("Document", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            date_para = doc.add_paragraph(
                f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M:%S UTC')}"
            )
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()

        # Split content into paragraphs and add to DOCX
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                # Clean up HTML-like tags if present
                para_text = para_text.replace("<p>", "").replace("</p>", "")
                para_text = para_text.replace("<br>", "\n").replace("<br/>", "\n")
                para_text = para_text.replace("&nbsp;", " ")
                
                doc.add_paragraph(para_text)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Save to storage
        storage_service = get_storage_service()
        # Generate file key using storage service
        file_key = storage_service.generate_file_key(
            organization_id=organization_id,
            deal_id=document_id,  # Use document_id as deal_id placeholder
            filename=f"{document_id}.docx",
            user_id=organization_id,  # Use organization_id as user_id placeholder
        )
        
        # Save file
        file_path = await storage_service.save_file(
            file_key=file_key,
            file_stream=buffer,
            organization_id=organization_id,
        )

        # Return file_key for retrieval, not full path
        return file_key

    @staticmethod
    async def export_to_html(
        content: str,
        organization_id: str,
        document_id: str,
        options: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        Export document content to HTML file.

        Args:
            content: Document content
            organization_id: Organization UUID
            document_id: Document UUID
            options: Export options

        Returns:
            File path where HTML is stored

        Raises:
            IOError: If file generation fails
        """
        options = options or {}
        include_cover_page = options.get("include_cover_page", False)

        # Create HTML document
        html_content = "<!DOCTYPE html>\n<html>\n<head>\n"
        html_content += "<meta charset='UTF-8'>\n"
        html_content += "<title>Document</title>\n"
        html_content += "<style>body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }</style>\n"
        html_content += "</head>\n<body>\n"

        if include_cover_page:
            html_content += f"<h1>Document</h1>\n"
            html_content += f"<p>Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M:%S UTC')}</p>\n"
            html_content += "<hr>\n"

        # Add content (escape HTML if needed)
        html_content += content.replace("\n", "<br>\n")

        html_content += "\n</body>\n</html>"

        # Convert to bytes
        html_bytes = html_content.encode("utf-8")
        file_stream = io.BytesIO(html_bytes)

        # Save to storage
        storage_service = get_storage_service()
        # Generate file key using storage service
        file_key = storage_service.generate_file_key(
            organization_id=organization_id,
            deal_id=document_id,  # Use document_id as deal_id placeholder
            filename=f"{document_id}.html",
            user_id=organization_id,  # Use organization_id as user_id placeholder
        )
        
        # Save file
        file_path = await storage_service.save_file(
            file_key=file_key,
            file_stream=file_stream,
            organization_id=organization_id,
        )

        return file_path

